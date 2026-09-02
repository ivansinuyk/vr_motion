"""Build the HAIT-formatted second-article DOCX from the accepted v5 template.

Rather than re-implementing the HAIT formatter (whose recipe is documented in
``second_article_plan_and_prompt.md``), this script clones the already
format-accepted first-article DOCX and edits it *in place*:

* the first-page single-column block (title, abstract, keywords, citation) is
  rewritten for article 2 while keeping authors, affiliations, DOI/UDC and the
  copyright line unchanged;
* the two-column body between INTRODUCTION and the section break before
  REFERENCES is fully replaced with the article-2 text, tables, and figures;
* every table is placed in a one-column island created with continuous section
  breaks, so wide tables span the full text width instead of being crushed into
  an 8.2 cm column (master-review item M29);
* the Ukrainian metadata block (title, abstract, keywords, citation) is
  rewritten;
* the mandatory AI-use disclosure is placed with the Acknowledgements, ahead of
  REFERENCES, per current HAIT policy;
* the REFERENCES list and, critically, the ABOUT THE AUTHORS photo table are
  left untouched so the embedded author photos, headers/footers, odd/even
  header-footer settings, margins, and column/section structure are preserved.

``article_package/second_article_manuscript.md`` is the single content source;
all numbers come from the analysis directory passed with ``--analysis-dir``.

Run:
    python build_second_article_docx.py \
        --manuscript article_package/second_article_manuscript.md \
        --analysis-dir second_article_outputs/v3 \
        --output "article_package/Стаття_Аспірант_Синюк_HAIT_article2_v4.docx"

Publication status: ``article_package/research_publication_status.md``
"""

from __future__ import annotations

import argparse
import copy
import glob
import re
import shutil
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

DEFAULT_TEMPLATE_GLOB = "article_package/*final_v5.docx"
DEFAULT_MANUSCRIPT = "article_package/second_article_manuscript.md"
DEFAULT_ANALYSIS_DIR = "second_article_outputs/v3"
DEFAULT_AGREEMENT_DIR = "second_article_outputs/annotation_agreement"
DEFAULT_OUTPUT = "article_package/Стаття_Аспірант_Синюк_HAIT_article2_v4.docx"

FONT = "Times New Roman"
COL_WIDTH_CM = 8.2
FULL_WIDTH_CM = 16.9
INDENT_CM = 0.75
TABLE_FONT_PT = 9


# --------------------------------------------------------------------------- #
# low-level run / paragraph helpers
# --------------------------------------------------------------------------- #
def _set_run_font(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


def rebuild_runs(par, specs):
    """Clear a paragraph's runs and add (text, bold, italic, size) specs."""
    for r in list(par.runs):
        r._element.getparent().remove(r._element)
    for text, bold, italic, size in specs:
        run = par.add_run(text)
        _set_run_font(run, size, bold, italic)


def set_text(par, text, size=None):
    if not par.runs:
        run = par.add_run(text)
        if size:
            _set_run_font(run, size)
        return
    par.runs[0].text = text
    for extra in list(par.runs[1:]):
        extra._element.getparent().remove(extra._element)


def find_par(doc, predicate):
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    raise LookupError("paragraph not found")


def find_par_after(doc, start_par, predicate):
    """Like ``find_par`` but skips everything up to and including ``start_par``.

    The body now carries its own AI-use disclosure next to the Acknowledgements,
    so a document-order search would match that copy instead of the template's
    declarations block.
    """
    seen = False
    for p in doc.paragraphs:
        if p._p is start_par._p:
            seen = True
            continue
        if seen and predicate(p.text):
            return p
    raise LookupError("paragraph not found after anchor")


def para_after_heading(doc, heading_text):
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == heading_text:
            for q in paras[i + 1:]:
                if q.text.strip():
                    return q
    raise LookupError(f"no paragraph after heading {heading_text!r}")


def remove_paragraph(par):
    par._element.getparent().remove(par._element)


# --------------------------------------------------------------------------- #
# body-block insertion (everything is inserted immediately before ``anchor``)
# --------------------------------------------------------------------------- #
def add_section_heading(anchor, text):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
    run = p.add_run(text.upper())
    _set_run_font(run, 11, bold=True)
    return p


def add_subheading(anchor, text):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(INDENT_CM)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
    run = p.add_run(text)
    _set_run_font(run, 11, bold=True, italic=True)
    return p


def add_body(anchor, text):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(INDENT_CM)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    _set_run_font(run, 11)
    return p


def add_numbered(anchor, items):
    for i, item in enumerate(items, start=1):
        p = anchor.insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format
        pf.first_line_indent = Cm(INDENT_CM)
        pf.line_spacing = 1.0
        run = p.add_run(f"{i}. {item}")
        _set_run_font(run, 11)


def add_figure(anchor, image_path, caption_num, caption_title, source="compiled by the authors"):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(COL_WIDTH_CM))

    cap = anchor.insert_paragraph_before()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.keep_with_next = True
    r1 = cap.add_run(f"Fig. {caption_num}. ")
    _set_run_font(r1, 11, italic=True)
    r2 = cap.add_run(caption_title)
    _set_run_font(r2, 11, bold=True)

    src = anchor.insert_paragraph_before()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.first_line_indent = Cm(0)
    src.paragraph_format.space_after = Pt(6)
    rs = src.add_run(f"Source: {source}")
    _set_run_font(rs, 8, bold=True, italic=True)


# --------------------------------------------------------------------------- #
# full-width table islands
# --------------------------------------------------------------------------- #
def _column_break_paragraph(anchor, body_sect, columns):
    """Insert a near-invisible paragraph whose section properties end a region.

    Word stores section properties on the *last* paragraph of a section, so a
    paragraph carrying ``cols/num=2`` closes the two-column run of body text and
    a paragraph carrying ``cols/num=1`` closes the full-width island that
    follows it.
    """
    p = anchor.insert_paragraph_before()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

    pPr = p._p.get_or_add_pPr()
    rpr = pPr.makeelement(qn("w:rPr"), {})
    size = pPr.makeelement(qn("w:sz"), {qn("w:val"): "2"})
    rpr.append(size)
    pPr.append(rpr)

    sect = copy.deepcopy(body_sect)
    type_el = sect.find(qn("w:type"))
    if type_el is None:
        type_el = sect.makeelement(qn("w:type"), {})
        sect.insert(0, type_el)
    type_el.set(qn("w:val"), "continuous")

    cols = sect.find(qn("w:cols"))
    if cols is None:
        cols = sect.makeelement(qn("w:cols"), {})
        sect.append(cols)
    cols.set(qn("w:num"), str(columns))

    pPr.append(sect)
    return p


def _repeat_header_row(table):
    """Mark the first row as a header so it repeats when the table breaks."""
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.makeelement(qn("w:tblHeader"), {})
    tr_pr.append(tbl_header)


def _style_cell(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    _set_run_font(run, TABLE_FONT_PT, bold=bold)


def add_table(doc, anchor, body_sect, table_num, title, df, widths, source="compiled by the authors"):
    """Render one table full-width between continuous section breaks."""
    _column_break_paragraph(anchor, body_sect, 2)

    tp = anchor.insert_paragraph_before()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.first_line_indent = Cm(0)
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.keep_with_next = True
    r1 = tp.add_run(f"Table {table_num}. ")
    _set_run_font(r1, 11, italic=True)
    r2 = tp.add_run(title)
    _set_run_font(r2, 11, bold=True)

    n_cols = len(df.columns)
    if widths is None or len(widths) != n_cols:
        widths = [FULL_WIDTH_CM / n_cols] * n_cols
    col_widths = [Cm(value) for value in widths]

    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        _style_cell(hdr[j], str(col), bold=True)
    _repeat_header_row(table)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            _style_cell(cells[j], str(row[col]))
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
    for j, column in enumerate(table.columns):
        column.width = col_widths[j]
    anchor._p.addprevious(table._tbl)

    src = anchor.insert_paragraph_before()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.first_line_indent = Cm(0)
    src.paragraph_format.space_after = Pt(6)
    rs = src.add_run(f"Source: {source}")
    _set_run_font(rs, 8, bold=True, italic=True)

    _column_break_paragraph(anchor, body_sect, 1)


# --------------------------------------------------------------------------- #
# manuscript parsing
# --------------------------------------------------------------------------- #
def _paragraphs(lines):
    paragraphs = []
    current = []
    for raw in lines:
        line = raw.strip().replace("`", "")
        if not line:
            if current:
                paragraphs.append(" ".join(current))
                current = []
            continue
        if re.match(r"^\d+\.\s+", line):
            if current:
                paragraphs.append(" ".join(current))
                current = []
            paragraphs.append(line)
        elif not line.startswith("---"):
            current.append(line)
    if current:
        paragraphs.append(" ".join(current))
    return paragraphs


def _parse_manuscript(path):
    """Parse the canonical Markdown source into metadata and body sections."""
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    metadata = {}
    current = None
    buffer = []
    body_start = None
    for idx, line in enumerate(lines):
        if line.strip() == "# BODY":
            if current:
                metadata[current] = _paragraphs(buffer)
            body_start = idx + 1
            break
        if line.startswith("## "):
            if current:
                metadata[current] = _paragraphs(buffer)
            current = line[3:].strip()
            buffer = []
        elif current:
            buffer.append(line)
    if body_start is None:
        raise ValueError(f"# BODY marker not found in {path}")

    body = {}
    section = None
    subsection = None
    buffer = []

    def flush():
        nonlocal buffer
        if section is None:
            buffer = []
            return
        paras = _paragraphs(buffer)
        target = body.setdefault(section, {"lead": [], "subsections": {}})
        if subsection is None:
            target["lead"].extend(paras)
        else:
            target["subsections"].setdefault(subsection, []).extend(paras)
        buffer = []

    for line in lines[body_start:]:
        if line.startswith("## "):
            flush()
            section = line[3:].strip()
            subsection = None
        elif line.startswith("### "):
            flush()
            subsection = line[4:].strip()
        else:
            buffer.append(line)
    flush()
    return metadata, body


@dataclass
class Content:
    title: str
    abstract: str
    keywords: str
    ua_title: str
    ua_abstract: str
    ua_keywords: str
    citation_en: str
    citation_ua: str
    intro: list
    litreview: list
    aim: str
    objectives: list
    methods: dict
    results: dict
    discussion: list
    conclusions: list
    acknowledgments: str


def load_content(manuscript_path) -> Content:
    meta, body = _parse_manuscript(manuscript_path)

    def one(name):
        values = meta.get(name, [])
        if not values:
            raise KeyError(f"Missing manuscript metadata section: {name}")
        return " ".join(values)

    aim_paras = body["RESEARCH AIM AND OBJECTIVES"]["lead"]
    return Content(
        title=one("TITLE"),
        abstract=one("ABSTRACT"),
        keywords=one("KEYWORDS"),
        ua_title=one("Ukrainian title"),
        ua_abstract=one("Ukrainian abstract"),
        ua_keywords=one("Ukrainian keywords"),
        citation_en=re.sub(r"^For citation:\s*", "", one("For citation (English)")),
        citation_ua=re.sub(r"^Для цитування:\s*", "", one("For citation (Ukrainian)")),
        intro=body["INTRODUCTION"]["lead"],
        litreview=body["LITERATURE REVIEW AND PROBLEM STATEMENT"]["lead"],
        aim=aim_paras[0],
        objectives=[
            re.sub(r"^\d+\.\s*", "", paragraph)
            for paragraph in aim_paras
            if re.match(r"^\d+\.\s+", paragraph)
        ],
        methods=body["MATERIALS AND METHODS"]["subsections"],
        results=body["RESEARCH RESULTS"]["subsections"],
        discussion=body["DISCUSSION OF RESULTS"]["lead"],
        conclusions=body["CONCLUSIONS"]["lead"],
        acknowledgments=" ".join(body["ACKNOWLEDGMENTS"]["lead"]),
    )


# --------------------------------------------------------------------------- #
# tables and figures
# --------------------------------------------------------------------------- #
TABLE_TITLES = {
    1: "Dataset and reference-subset characteristics",
    2: "Operational definitions of manually annotated events",
    3: "Annotation agreement between and within annotators",
    4: "Event-frame agreement with the consensus reference",
    5: "Clubhead localization agreement with the consensus by swing phase",
    6: "Perturbation response and rank preservation of exported metrics",
    7: "Nested ablation of actual production stages",
}

TABLE_WIDTHS = {
    1: [4.5, 6.2, 6.2],
    2: [3.6, 7.7, 5.6],
    3: [4.4, 0.9, 2.7, 2.2, 1.7, 1.9, 3.1],
    4: [4.4, 1.0, 4.2, 7.3],
    5: [3.6, 2.4, 5.4, 5.5],
    6: [3.6, 2.4, 1.8, 1.8, 2.8, 4.5],
    7: [5.4, 3.2, 3.6, 4.7],
}

FIGURE_TITLES = {
    1: ("fig_study_design.png", "Study design and evidence scope"),
    2: ("fig_annotated_frame.png", "Operational event and selected-frame clubhead annotations"),
    3: ("fig_annotation_agreement.png", "Human annotation agreement between and within annotators"),
    4: ("fig_event_frame_identity.png", "Automatic versus manual consensus event frames"),
    5: ("fig_trajectory_error_by_phase.png", "Session-level clubhead agreement by swing phase"),
    6: ("fig_sensitivity_metric_heatmap.png", "Metric response by perturbation family"),
    7: ("fig_ablation_jerk_reduction.png", "Root-mean-square jerk across production stages"),
}

GENERATED_TABLE_FILES = {
    1: "table1_dataset_subset.csv",
    2: "table2_annotation_protocol.csv",
    4: "table3_event_timing.csv",
    5: "table4_trajectory_error.csv",
    7: "table6_ablation.csv",
}


HEADER_REWRITES = {"RMS jerk, m/s^3": "RMS jerk, m/s\u00b3"}


def load_generated_table(tables_dir: Path, table_num: int) -> pd.DataFrame:
    df = pd.read_csv(tables_dir / GENERATED_TABLE_FILES[table_num])
    return df.rename(columns=HEADER_REWRITES)


def _fmt(value, digits=1):
    if pd.isna(value):
        return "-"
    return f"{float(value):.{digits}f}".rstrip("0").rstrip(".") if digits else f"{float(value):.0f}"


def build_agreement_table(agreement_dir: Path) -> pd.DataFrame:
    """Table 3: human-human agreement, straight from the agreement summary."""
    summary = pd.read_csv(agreement_dir / "annotation_agreement_summary.csv")

    event_labels = {
        "address": "Address",
        "top_backswing": "Top of backswing",
        "downswing_transition": "Downswing transition",
        "impact": "Impact",
        "ALL": "All events",
    }
    plan = [
        ("inter_rater", "Between annotators", ["address", "top_backswing", "downswing_transition", "impact", "ALL"]),
        ("intra_rater_ivan", "Within annotator, repeat round", ["ALL"]),
    ]

    rows = []
    for comparison, comparison_label, groups in plan:
        block = summary[summary["comparison"] == comparison]
        for group in groups:
            row = block[(block["domain"] == "event") & (block["group"] == group)]
            if row.empty:
                continue
            r = row.iloc[0]
            median = r["session_median_abs_frame"] if pd.notna(r["session_median_abs_frame"]) else r["median_abs_frame"]
            rows.append(
                {
                    "Comparison and measure": f"{comparison_label}: {event_labels[group].lower()}, frames",
                    "n": int(r["n_pairs"]),
                    "Median (95% CI)": (
                        f"{_fmt(median)} "
                        f"({_fmt(r['session_median_abs_frame_ci_low'])}-"
                        f"{_fmt(r['session_median_abs_frame_ci_high'])})"
                    ),
                    "P95 / max": f"{_fmt(r['p95_abs_frame'])} / {_fmt(r['max_abs_frame'])}",
                    "Exact, %": _fmt(100 * r["exact_frame_frac"], 0),
                    "Within 2, %": _fmt(100 * r["within_2_frame_frac"], 0),
                    "Bias (95% LoA)": (
                        f"{_fmt(r['bias_frame'], 2)} "
                        f"({_fmt(r['loa_low_frame'])} to {_fmt(r['loa_high_frame'])})"
                    ),
                }
            )

    for comparison, comparison_label, _ in plan:
        row = summary[(summary["comparison"] == comparison) & (summary["domain"] == "point")]
        if row.empty:
            continue
        r = row.iloc[0]
        median = r["session_median_pixel"] if pd.notna(r["session_median_pixel"]) else r["median_pixel"]
        rows.append(
            {
                "Comparison and measure": f"{comparison_label}: clubhead clicks, px",
                "n": int(r["n_pairs"]),
                "Median (95% CI)": (
                    f"{_fmt(median)} "
                    f"({_fmt(r['session_median_pixel_ci_low'])}-{_fmt(r['session_median_pixel_ci_high'])})"
                ),
                "P95 / max": f"{_fmt(r['p95_pixel'])} / {_fmt(r['max_pixel'])}",
                "Exact, %": "-",
                "Within 2, %": "-",
                "Bias (95% LoA)": f"{int(r['n_over_50px'])} pairs > 50 px",
            }
        )
    return pd.DataFrame(rows)


def build_sensitivity_table(tables_dir: Path) -> pd.DataFrame:
    """Table 6: the two generated sensitivity tables joined on metric name."""
    sens = pd.read_csv(tables_dir / "table5_sensitivity.csv")
    rank = pd.read_csv(tables_dir / "table7_robustness_ranking.csv")
    merged = rank.merge(sens, on="Metric", how="left", validate="one_to_one")
    return pd.DataFrame(
        {
            "Metric": merged["Metric"],
            "Median change, %": merged["Median Δsym, %"],
            "Worst, %": merged["Worst, %"],
            "Median ρ": merged["Median ρ"],
            "Low/mod./high": merged["Scenarios low/moderate/high"],
            "Interpretation": merged["Interpretation"],
        }
    )


# --------------------------------------------------------------------------- #
# document assembly
# --------------------------------------------------------------------------- #
def rewrite_first_page(doc, content: Content):
    title_p = find_par(doc, lambda t: t.startswith("Markerless video-based golf-stick motion analysis using Kalman"))
    set_text(title_p, content.title)

    abstract_p = para_after_heading(doc, "ABSTRACT")
    set_text(abstract_p, content.abstract)

    kw_p = find_par(doc, lambda t: t.startswith("Keywords:"))
    rebuild_runs(kw_p, [("Keywords: ", True, False, 9), (content.keywords, False, False, 9)])

    cite_p = find_par(doc, lambda t: t.startswith("For citation:"))
    rebuild_runs(cite_p, [("For citation: ", True, True, 8), (content.citation_en, True, False, 8)])


def rewrite_ukrainian(doc, content: Content):
    ua_title = find_par(doc, lambda t: t.startswith("Безмаркерний відеоаналіз руху ключки"))
    set_text(ua_title, content.ua_title)

    ua_abs = para_after_heading(doc, "АНОТАЦІЯ")
    set_text(ua_abs, content.ua_abstract)

    ua_kw = find_par(doc, lambda t: t.startswith("Ключові слова:"))
    rebuild_runs(ua_kw, [("Ключові слова: ", True, False, 9), (content.ua_keywords, False, False, 9)])

    ua_cite = find_par(doc, lambda t: t.startswith("Для цитування:"))
    rebuild_runs(ua_cite, [("Для цитування: ", True, True, 8), (content.citation_ua, True, False, 8)])


def find_body_anchor(doc):
    paras = doc.paragraphs
    intro = None
    for i, p in enumerate(paras):
        if p.text.strip() == "INTRODUCTION":
            intro = p
            intro_i = i
            break
    if intro is None:
        raise LookupError("INTRODUCTION not found")
    for p in paras[intro_i + 1:]:
        if p._p.find(qn("w:pPr")) is not None and p._p.find(qn("w:pPr")).find(qn("w:sectPr")) is not None:
            return intro, p
    raise LookupError("section break after body not found")


def remove_old_body(doc, intro_p, anchor_p):
    body = doc.element.body
    removing = False
    to_remove = []
    for child in list(body.iterchildren()):
        if child is intro_p._p:
            removing = True
        if child is anchor_p._p:
            break
        if removing:
            to_remove.append(child)
    for el in to_remove:
        body.remove(el)


AI_DISCLOSURE = (
    "Use of generative AI tools: AI-assisted language revision, consistency checking, and code "
    "review were performed with GPT-5.6 Sol, Claude Opus 5, and Cursor Grok 4.6 in Cursor during "
    "manuscript preparation in 2026. The authors reviewed the resulting text and remain responsible "
    "for the data, methods, interpretation, references, and final manuscript."
)


def build_body(doc, anchor, body_sect, content: Content, fig_dir: Path, tables_dir: Path, agreement_dir: Path):
    def figure(number, source="compiled by the authors"):
        name, caption = FIGURE_TITLES[number]
        add_figure(anchor, fig_dir / name, number, caption, source=source)

    def table(number, df, source="compiled by the authors"):
        add_table(
            doc, anchor, body_sect, number, TABLE_TITLES[number], df, TABLE_WIDTHS.get(number), source=source
        )

    add_section_heading(anchor, "INTRODUCTION")
    for text in content.intro:
        add_body(anchor, text)
    figure(1)

    add_section_heading(anchor, "LITERATURE REVIEW AND PROBLEM STATEMENT")
    for text in content.litreview:
        add_body(anchor, text)

    add_section_heading(anchor, "RESEARCH AIM AND OBJECTIVES")
    add_body(anchor, content.aim)
    add_body(anchor, "The research objectives are:")
    add_numbered(anchor, content.objectives)

    add_section_heading(anchor, "MATERIALS AND METHODS")
    for sub, texts in content.methods.items():
        add_subheading(anchor, sub)
        for text in texts:
            add_body(anchor, text)
        if sub.startswith("Dataset, sampling"):
            table(1, load_generated_table(tables_dir, 1))
        elif sub.startswith("Manual annotation protocol"):
            table(2, load_generated_table(tables_dir, 2))
            figure(2)

    add_section_heading(anchor, "RESEARCH RESULTS")
    for sub, paras in content.results.items():
        add_subheading(anchor, sub)
        for text in paras:
            add_body(anchor, text)
        if sub.startswith("Annotation agreement"):
            table(
                3,
                build_agreement_table(agreement_dir),
                source=(
                    "compiled by the authors; human annotation only; "
                    "confidence intervals resample sessions"
                ),
            )
            figure(3)
        elif sub.startswith("Time-base audit"):
            table(4, load_generated_table(tables_dir, 4))
            figure(4)
        elif sub.startswith("Clubhead localization agreement"):
            table(
                5,
                load_generated_table(tables_dir, 5),
                source="compiled by the authors; n_s - sessions; n_p - points",
            )
            figure(5)
        elif sub.startswith("Perturbation sensitivity"):
            table(
                6,
                build_sensitivity_table(tables_dir),
                source=(
                    "compiled by the authors; response bands are study-specific descriptive "
                    "labels, not measurement tolerances"
                ),
            )
            figure(6)
        elif sub.startswith("Production-stage ablation"):
            table(7, load_generated_table(tables_dir, 7))
            figure(7)

    add_section_heading(anchor, "DISCUSSION OF RESULTS")
    for text in content.discussion:
        add_body(anchor, text)

    add_section_heading(anchor, "CONCLUSIONS")
    for text in content.conclusions:
        add_body(anchor, text)

    add_section_heading(anchor, "ACKNOWLEDGMENTS")
    add_body(anchor, content.acknowledgments)
    add_body(anchor, AI_DISCLOSURE)


# --------------------------------------------------------------------------- #
# references and declarations
# --------------------------------------------------------------------------- #
def _reference_specs(number, authors, title, journal, details):
    return [
        (f"{number}. {authors} “{title}”. ", False, False, 11),
        (journal, False, True, 11),
        (f". {details}", False, False, 11),
    ]


def _format_reference_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(INDENT_CM)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def rewrite_references_and_declarations(doc):
    replacements = {
        "7. Challis": _reference_specs(
            7,
            "Davis, D. J. & Challis, J. H.",
            "Automatic segment filtering procedure for processing non-stationary signals",
            "Journal of Biomechanics",
            "2020; Vol. 101: 109619. DOI: https://doi.org/10.1016/j.jbiomech.2020.109619.",
        ),
        "24. Godfrey": _reference_specs(
            24,
            "Kim, S. E., Burket Koltsov, J. C., Richards, A. W., Zhou, J., "
            "Schadl, K., Ladd, A. L. & Rose, J.",
            "Validation of inertial measurement units for analyzing golf swing rotational biomechanics",
            "Sensors",
            "2023; Vol. 23, No. 20: 8433. DOI: https://doi.org/10.3390/s23208433.",
        ),
        "30. Sweeting": _reference_specs(
            30,
            "Ingwersen, C. K., Nørtoft Jensen, J., Rieger Hannemose, M. & "
            "Bjorholm Dahl, A.",
            "Evaluating current state of monocular 3D pose models for golf",
            "Proceedings of the Northern Lights Deep Learning Workshop",
            "2023; Vol. 4. DOI: https://doi.org/10.7557/18.6793.",
        ),
    }
    for prefix, specs in replacements.items():
        paragraph = find_par(doc, lambda text, p=prefix: text.startswith(p))
        rebuild_runs(paragraph, specs)
        _format_reference_paragraph(paragraph)

    conflict = find_par(doc, lambda text: text.startswith("Conflicts of Interest:"))
    new_references = [
        _reference_specs(
            31,
            "Yamamoto, K., Hasegawa, Y., Suzuki, T., Suzuki, H., Tanabe, H. & Fujii, K.",
            "Extracting proficiency differences and individual characteristics in golfers' "
            "swing using single-video markerless motion analysis",
            "Frontiers in Sports and Active Living",
            "2023; Vol. 5: 1272038. DOI: https://doi.org/10.3389/fspor.2023.1272038.",
        ),
        _reference_specs(
            32,
            "Syniuk, I. M., Maksymov, M. V., Maksymov, O. M. & Iyer, K.",
            "Markerless video-based golf-stick motion analysis using Kalman filtering "
            "and RTS smoothing",
            "Unpublished manuscript",
            "2026.",
        ),
    ]
    for specs in new_references:
        paragraph = conflict.insert_paragraph_before()
        rebuild_runs(paragraph, specs)
        _format_reference_paragraph(paragraph)

    # HAIT policy places the AI-use disclosure with the Acknowledgements, ahead
    # of REFERENCES; build_body writes it there, so drop the template copy that
    # sits in the post-references declarations block.
    template_ai = find_par_after(doc, conflict, lambda text: text.startswith("Use of generative AI tools:"))

    def insert_statement(text):
        paragraph = template_ai.insert_paragraph_before()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(text)
        _set_run_font(run, 11)

    insert_statement(
        "Ethics and consent: No ethics-committee approval was required. The analysed "
        "recordings are anonymized public demonstration footage used for algorithmic "
        "evaluation; they do not constitute identifiable human-participant research data "
        "for the purposes reported here. No participant image is reproduced."
    )
    insert_statement(
        "Data availability: Analysis code, derived aggregate tables, and de-identified "
        "session-level outputs can be made available on request. Source demonstration "
        "videos are not separately released as a public dataset in this submission."
    )
    remove_paragraph(template_ai)


def remove_missing_photo_placeholders(doc):
    """Clear literal photo prompts without rebuilding the author table."""
    if not doc.tables:
        return
    author_table = doc.tables[-1]
    for row in author_table.rows:
        photo_cell = row.cells[0]
        if photo_cell.text.strip().startswith("[3×4 cm color photo]"):
            # These cells contain no drawing; editing them in place preserves every
            # embedded photograph in the other author rows.
            if not photo_cell._tc.xpath(".//w:drawing"):
                photo_cell.text = ""


def main():
    parser = argparse.ArgumentParser(description="Build the article-2 HAIT DOCX.")
    parser.add_argument("--template", default=None, help=f"defaults to {DEFAULT_TEMPLATE_GLOB}")
    parser.add_argument("--manuscript", default=DEFAULT_MANUSCRIPT)
    parser.add_argument("--analysis-dir", default=DEFAULT_ANALYSIS_DIR)
    parser.add_argument("--agreement-dir", default=DEFAULT_AGREEMENT_DIR)
    parser.add_argument("--output", default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    template = args.template or glob.glob(DEFAULT_TEMPLATE_GLOB)[0]
    analysis_dir = Path(args.analysis_dir)
    fig_dir = analysis_dir / "figures"
    tables_dir = analysis_dir / "article_tables"
    agreement_dir = Path(args.agreement_dir)
    output = Path(args.output)

    if Path(template).resolve() == output.resolve():
        raise SystemExit("refusing to overwrite the template with the output")
    v2 = Path("article_package") / "Стаття_Аспірант_Синюк_HAIT_article2_v2.docx"
    if v2.exists() and output.resolve() == v2.resolve():
        raise SystemExit("refusing to overwrite the frozen v2 DOCX")

    content = load_content(args.manuscript)

    missing = [
        str(path)
        for path in (
            [tables_dir / name for name in GENERATED_TABLE_FILES.values()]
            + [tables_dir / "table5_sensitivity.csv", tables_dir / "table7_robustness_ranking.csv"]
            + [agreement_dir / "annotation_agreement_summary.csv"]
            + [fig_dir / name for name, _ in FIGURE_TITLES.values()]
        )
        if not path.exists()
    ]
    if missing:
        raise SystemExit("missing required build inputs:\n  " + "\n  ".join(missing))

    shutil.copyfile(template, output)
    doc = Document(str(output))

    rewrite_first_page(doc, content)
    intro_p, anchor_p = find_body_anchor(doc)
    body_sect = anchor_p._p.find(qn("w:pPr")).find(qn("w:sectPr"))
    remove_old_body(doc, intro_p, anchor_p)
    build_body(doc, anchor_p, body_sect, content, fig_dir, tables_dir, agreement_dir)
    rewrite_ukrainian(doc, content)
    rewrite_references_and_declarations(doc)
    remove_missing_photo_placeholders(doc)

    doc.save(str(output))
    print(f"saved {output}")


if __name__ == "__main__":
    main()
